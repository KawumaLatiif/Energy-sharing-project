"""
Generate USSD Features Report (Word) with screenshots.
Run from project root:
  backend\\venv\\Scripts\\python.exe scripts\\generate_ussd_report.py
"""
from __future__ import annotations

import json
import textwrap
from datetime import datetime
from pathlib import Path

import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = PROJECT_ROOT / "docs" / "ussd-report"
SCREENSHOTS_DIR = OUT_DIR / "screenshots"
DOCX_PATH = OUT_DIR / "USSD_Features_Report.docx"

USSD_URL = "http://127.0.0.1:8000/api/v1/ussd/entry/"
SIMULATOR_URL = "http://127.0.0.1:3000/ussd-simulator"
PHONE = "+256701234567"
SERVICE_CODE = "*123#"


def normalize_ussd_response(raw: str) -> str:
    text = raw.strip()
    if text.startswith('"') and text.endswith('"'):
        try:
            text = json.loads(text)
        except json.JSONDecodeError:
            text = text.strip('"')
    if text.startswith("CON "):
        return text[4:]
    if text.startswith("END "):
        return text[4:]
    return text


def ussd_call(session_id: str, text: str) -> str:
    payload = {
        "sessionId": session_id,
        "serviceCode": SERVICE_CODE,
        "phoneNumber": PHONE,
        "text": text,
    }
    try:
        resp = requests.post(USSD_URL, json=payload, timeout=15)
        resp.raise_for_status()
        return normalize_ussd_response(resp.text)
    except requests.RequestException as exc:
        return f"[USSD API unavailable: {exc}]"


def render_ussd_screen(title: str, body: str, path: Path, footer: str = "") -> bool:
    """Render a phone-style USSD screen PNG."""
    lines = [title, ""] + body.split("\n")
    if footer:
        lines.extend(["", footer])

    font = ImageFont.load_default()
    line_height = 18
    padding = 24
    width = 420
    height = max(320, padding * 2 + line_height * len(lines) + 40)

    img = Image.new("RGB", (width, height), color=(15, 76, 129))
    draw = ImageDraw.Draw(img)
    draw.rounded_rectangle(
        (12, 12, width - 12, height - 12),
        radius=16,
        fill=(245, 248, 252),
    )

    y = 28
    for i, line in enumerate(lines):
        color = (20, 20, 20) if i > 0 else (15, 76, 129)
        if i == 0:
            draw.text((28, y), line, fill=color, font=font)
        else:
            wrapped = textwrap.wrap(line, width=42) or [""]
            for wline in wrapped:
                draw.text((28, y), wline, fill=color, font=font)
                y += line_height
            continue
        y += line_height

    draw.text((28, height - 32), "gPawa USSD", fill=(100, 100, 100), font=font)
    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path)
    return True


def capture_api_screenshots() -> dict[str, Path]:
    """Capture representative USSD menu screens via API."""
    shots: dict[str, Path] = {}
    flows = [
        ("01_main_menu", "", "Main Menu"),
        ("02_wallet_meter", "1", "Wallet & Meter"),
        ("03_buy_menu", "2", "Buy Units Menu"),
        ("04_loans_menu", "3", "Loans Menu"),
        ("05_loan_latest", "3*1", "Latest Loan"),
        ("06_loan_stats", "3*5", "Loan Statistics"),
        ("07_share_menu", "4", "Share Units Menu"),
        ("08_tokens", "5", "My Tokens"),
        ("09_buy_amount_prompt", "2*1", "Buy Units — Amount"),
    ]

    for key, text, title in flows:
        sid = f"RPT-{key}"
        body = ussd_call(sid, text)
        out = SCREENSHOTS_DIR / f"{key}.png"
        render_ussd_screen(title, body, out, footer=f"text={text or '(empty)'}")
        shots[key] = out

    # Buy flow (may error if active loan — still useful for report)
    sid = "RPT-10_buy_result"
    body = ussd_call(sid, "2*1*30000")
    out = SCREENSHOTS_DIR / "10_buy_purchase_result.png"
    render_ussd_screen("Buy Units — Purchase", body, out, footer="text=2*1*30000")
    shots["10_buy_purchase_result"] = out

    # Share flow (receiver meter must exist in seed data)
    body = ussd_call("RPT-11_share", "4*1*1234567892*5")
    out = SCREENSHOTS_DIR / "11_share_initiate.png"
    if not out.exists():
        render_ussd_screen("Share Initiate", body, out, footer="text=4*1*1234567892*5")
    shots["11_share_initiate"] = out

    return shots


def capture_playwright_screenshots() -> dict[str, Path]:
    """Capture web simulator UI if Playwright + frontend are available."""
    shots: dict[str, Path] = {}
    try:
        from playwright.sync_api import sync_playwright
    except ImportError:
        return shots

    try:
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True)
            page = browser.new_page(viewport={"width": 1280, "height": 900})
            page.goto(SIMULATOR_URL, wait_until="networkidle", timeout=60000)

            # Initial UI
            p1 = SCREENSHOTS_DIR / "sim_01_simulator_initial.png"
            page.screenshot(path=str(p1), full_page=True)
            shots["sim_01_simulator_initial"] = p1

            page.get_by_role("button", name="Open Menu").click()
            page.wait_for_timeout(1500)
            p2 = SCREENSHOTS_DIR / "sim_02_main_menu_conversation.png"
            page.screenshot(path=str(p2), full_page=True)
            shots["sim_02_main_menu_conversation"] = p2

            # Wallet flow
            page.locator('input[placeholder*="Example"]').fill("1")
            page.get_by_role("button", name="Send Reply").click()
            page.wait_for_timeout(1500)
            p3 = SCREENSHOTS_DIR / "sim_03_wallet_meter_flow.png"
            page.screenshot(path=str(p3), full_page=True)
            shots["sim_03_wallet_meter_flow"] = p3

            page.get_by_role("button", name="New Session").click()
            page.wait_for_timeout(500)
            page.get_by_role("button", name="Open Menu").click()
            page.wait_for_timeout(1000)
            page.locator('input[placeholder*="Example"]').fill("3")
            page.get_by_role("button", name="Send Reply").click()
            page.wait_for_timeout(1000)
            page.locator('input[placeholder*="Example"]').fill("1")
            page.get_by_role("button", name="Send Reply").click()
            page.wait_for_timeout(1500)
            p4 = SCREENSHOTS_DIR / "sim_04_loans_latest_flow.png"
            page.screenshot(path=str(p4), full_page=True)
            shots["sim_04_loans_latest_flow"] = p4

            browser.close()
    except Exception as exc:
        print(f"Playwright capture skipped: {exc}")
    return shots


def add_heading(doc: Document, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p


def add_image_or_placeholder(
    doc: Document,
    image_path: Path | None,
    caption: str,
    placeholder_hint: str,
):
    doc.add_paragraph(caption, style="Caption")
    if image_path and image_path.exists():
        try:
            doc.add_picture(str(image_path), width=Inches(5.5))
        except Exception:
            add_para(
                doc,
                f"[Screenshot placeholder: {placeholder_hint} — file present but could not embed]",
                bold=True,
            )
    else:
        p = doc.add_paragraph()
        run = p.add_run(f"[Screenshot placeholder: {placeholder_hint}]")
        run.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)
        p.paragraph_format.space_after = Pt(12)
        box = doc.add_paragraph()
        box_run = box.add_run(
            "Insert screenshot here showing the described screen or simulator step."
        )
        box_run.italic = True
    doc.add_paragraph()


def build_document(api_shots: dict[str, Path], sim_shots: dict[str, Path]) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = title.add_run("gPawa\nUSSD Integration\nFeatures Report")
    t_run.bold = True
    t_run.font.size = Pt(26)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"Generated: {datetime.now().strftime('%d %B %Y')}\n")
    sub.add_run("Project: gPawa (gPawa)\n")
    sub.add_run(f"Test handset: {PHONE}")

    doc.add_page_break()

    add_heading(doc, "1. Executive Summary", 1)
    add_para(
        doc,
        "This report documents the USSD channel implemented for gPawa. "
        "Subscribers dial a short code (e.g. *123#) and navigate text menus to view wallet "
        "and meter information, buy electricity units, manage loans, share units with OTP "
        "verification, and list active meter tokens. The backend exposes a single entry "
        "point compatible with Africa's Talking, and a browser-based simulator is provided "
        "for local testing without a physical handset.",
    )

    add_heading(doc, "2. Architecture", 2)
    add_para(doc, "Components:", bold=True)
    for item in [
        "USSD provider (Africa's Talking) or web simulator sends sessionId, serviceCode, phoneNumber, text.",
        "Django endpoint POST /api/v1/ussd/entry/ resolves the user by phone number.",
        "ussd.UssdSession stores menu context, shortcuts, and deduplicated responses.",
        "Business logic reuses buy-units, loan scoring, share OTP, and token models from the main app.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_image_or_placeholder(
        doc,
        sim_shots.get("sim_01_simulator_initial"),
        "Figure 2.1 — USSD Web Simulator (initial screen)",
        "USSD Web Simulator at /ussd-simulator with session fields and action buttons",
    )

    add_heading(doc, "3. Request and Response Contract", 1)
    add_para(
        doc,
        "Each dial step POSTs JSON (or form fields) with cumulative menu path in text "
        "(segments joined by *). Responses are plain text prefixed with CON (continue) or END (terminate).",
    )
    doc.add_paragraph('Example: text "2*1*30000" = Buy Units → Start purchase → 30000 UGX.')

    add_heading(doc, "4. Main Menu", 1)
    add_para(doc, "Opening the session (empty text) returns:", bold=True)
    menu_lines = (
        "gPawa\n"
        "1. Wallet & Meter\n"
        "2. Buy Units\n"
        "3. Loans\n"
        "4. Share Units\n"
        "5. My Tokens\n"
        "6. Exit"
    )
    doc.add_paragraph(menu_lines)

    add_image_or_placeholder(
        doc,
        api_shots.get("01_main_menu"),
        "Figure 4.1 — Main menu (USSD screen)",
        "Main menu CON response on handset or simulator",
    )
    add_image_or_placeholder(
        doc,
        sim_shots.get("sim_02_main_menu_conversation"),
        "Figure 4.2 — Main menu in web simulator conversation panel",
        "Simulator after Open Menu showing CON response in history",
    )

    # Feature 1
    add_heading(doc, "5. Feature: Wallet & Meter (Option 1)", 1)
    add_para(
        doc,
        "Path: text = 1. One-step END response showing unit wallet balance, registered meter "
        "number (or Not registered), and total outstanding on disbursed loans (UGX).",
    )
    add_image_or_placeholder(
        doc,
        api_shots.get("02_wallet_meter"),
        "Figure 5.1 — Wallet & Meter response",
        "END screen with wallet balance, meter number, outstanding loan",
    )
    add_image_or_placeholder(
        doc,
        sim_shots.get("sim_03_wallet_meter_flow"),
        "Figure 5.2 — Wallet flow in simulator",
        "Simulator conversation after sending reply 1 from main menu",
    )

    # Feature 2
    add_heading(doc, "6. Feature: Buy Units (Option 2)", 1)
    add_para(doc, "Submenu (text = 2):", bold=True)
    doc.add_paragraph("1. Start purchase\n2. Check payment status", style="List Bullet")
    add_para(
        doc,
        "Start purchase (2*1*amount): requires registered meter; blocked if user has any loan "
        "not COMPLETED or REJECTED. Sandbox mode simulates MoMo payment in ~10 seconds and credits "
        "the unit wallet. Check status (2*2*txId) supports shortcut 0 for last transaction in session.",
    )
    add_image_or_placeholder(
        doc,
        api_shots.get("03_buy_menu"),
        "Figure 6.1 — Buy Units submenu",
        "CON Buy Units submenu",
    )
    add_image_or_placeholder(
        doc,
        api_shots.get("09_buy_amount_prompt"),
        "Figure 6.2 — Enter amount prompt",
        "CON Enter amount in UGX after 2*1",
    )
    add_image_or_placeholder(
        doc,
        api_shots.get("10_buy_purchase_result"),
        "Figure 6.3 — Purchase initiation result",
        "END response with TxID and PENDING/SUCCESS or error message",
    )

    # Feature 3
    add_heading(doc, "7. Feature: Loans (Option 3)", 1)
    add_para(doc, "Submenu (text = 3):", bold=True)
    for item in [
        "1. Latest loan — summary of most recent application",
        "2. Apply loan — amount 5000–200000 UGX, credit scoring, tier approval, and automatic disbursement",
        "3. Repay loan — disbursed loans; LoanID then amount",
        "4. Loan stats — pending/active counts and outstanding total",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_image_or_placeholder(
        doc,
        api_shots.get("04_loans_menu"),
        "Figure 7.1 — Loans submenu",
        "CON Loans menu with five options",
    )
    add_image_or_placeholder(
        doc,
        api_shots.get("05_loan_latest"),
        "Figure 7.2 — Latest loan details",
        "END loan ID, ref, status, amounts, outstanding",
    )
    add_image_or_placeholder(
        doc,
        api_shots.get("06_loan_stats"),
        "Figure 7.3 — Loan statistics",
        "END pending/active/outstanding summary",
    )
    add_image_or_placeholder(
        doc,
        sim_shots.get("sim_04_loans_latest_flow"),
        "Figure 7.4 — Loans flow in simulator (3 → 1)",
        "Simulator showing loans submenu and latest loan END",
    )

    # Feature 4
    add_heading(doc, "8. Feature: Share Units (Option 4)", 1)
    add_para(
        doc,
        "Submenu: 1 Initiate share (meter → units, min 2), 2 Verify OTP (ref + 6-digit code). "
        "Initiate creates a PENDING ShareTransaction and verification code (purpose share_units). "
        "Web flow emails OTP via Celery; USSD initiate currently stores OTP in DB only — verify "
        "using admin/DB during testing unless email hook is added. Ref shortcut 0 reuses last share ref.",
    )
    add_image_or_placeholder(
        doc,
        api_shots.get("07_share_menu"),
        "Figure 8.1 — Share Units submenu",
        "CON Share Units with Initiate and Verify OTP",
    )
    share_init = api_shots.get("11_share_initiate")
    if share_init and share_init.exists():
        add_image_or_placeholder(
            doc,
            share_init,
            "Figure 8.2 — Share initiate completed (END with transaction ref)",
            "END after 4*1*meter*units showing Ref SHARE-...",
        )
    else:
        add_image_or_placeholder(
            doc,
            None,
            "Figure 8.2 — Share initiate completed (END with transaction ref)",
            "END after 4*1*meter*units showing Ref SHARE-...",
        )

    share_verify = SCREENSHOTS_DIR / "12_share_verify_success.png"
    add_image_or_placeholder(
        doc,
        share_verify if share_verify.exists() else None,
        "Figure 8.3 — Share OTP verification success",
        "END after 4*2*ref*otp showing units transferred or token issued",
    )

    # Feature 5
    add_heading(doc, "9. Feature: My Tokens (Option 5)", 1)
    add_para(
        doc,
        "Lists up to three unused meter tokens (token | units | source). END with message if none.",
    )
    add_image_or_placeholder(
        doc,
        api_shots.get("08_tokens"),
        "Figure 9.1 — Active tokens list",
        "END Active tokens listing",
    )

    add_heading(doc, "10. Session Management", 1)
    add_para(
        doc,
        "UssdSession persists sessionId for 15 minutes, stores context (last buy TxID, last share ref), "
        "and returns cached last_response when the provider retries the same text — preventing duplicate charges.",
    )

    add_heading(doc, "11. Features Not on USSD", 1)
    for item in [
        "Meter transfer (available on web/API only)",
        "Account registration and login",
        "Admin operations",
        "Dedicated USSD MoMo loan repayment channel (web has repay/momo/)",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_heading(doc, "12. Testing and Deployment", 1)
    add_para(
        doc,
        "Local: Django :8000, Next.js :3000, optional PostgreSQL seed (sample_full_dump_heavy.sql). "
        "Production: configure Africa's Talking callback to https://<host>/api/v1/ussd/entry/. "
        "Test MSISDN must exist in accounts_user with last-9-digit phone matching.",
    )
    add_image_or_placeholder(
        doc,
        None,
        "Figure 12.1 — Dashboard navigation to USSD Simulator",
        "Sidebar or account menu link to /ussd-simulator",
    )
    add_image_or_placeholder(
        doc,
        None,
        "Figure 12.2 — Africa's Talking channel configuration",
        "AT dashboard showing callback URL and service code",
    )

    add_heading(doc, "13. Source Reference", 1)
    refs = [
        "backend/ussd/views.py — menu handler",
        "backend/ussd/models.py — UssdSession",
        "frontend/src/app/ussd-simulator/page.tsx — browser simulator",
        "USSD_INTEGRATION.md — technical integration guide",
    ]
    for ref in refs:
        doc.add_paragraph(ref, style="List Bullet")

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)
    print(f"Report saved: {DOCX_PATH}")


def main():
    print("Capturing API-based USSD screens...")
    api_shots = capture_api_screenshots()
    print("Capturing web simulator screenshots (Playwright)...")
    sim_shots = capture_playwright_screenshots()
    print("Building Word document...")
    build_document(api_shots, sim_shots)
    print("Done.")


if __name__ == "__main__":
    main()
